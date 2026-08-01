from __future__ import annotations

from pathlib import Path
from datetime import date

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
ASSET_DIR = OUT_DIR / "assets_ieee1016"
OUTPUT = OUT_DIR / "SDD_IEEE1016_RAG_Ley19628.docx"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
WHITE = "FFFFFF"
BLACK = "000000"
GREEN = "287A4B"
AMBER = "9A6700"
RED = "9B1C1C"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def font_run(run, size=11, bold=False, italic=False, color=BLACK, name="Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction: str, fallback: str = "") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instr, separate, text, end])
    font_run(run, size=9, color=MID_GRAY)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MID_GRAY)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(7)


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.82)
        section.bottom_margin = Inches(0.78)
        section.left_margin = Inches(0.86)
        section.right_margin = Inches(0.86)
        section.header_distance = Inches(0.38)
        section.footer_distance = Inches(0.4)
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("SDD · RAG Ley 19.628  |  IEEE 1016-2009")
        font_run(run, size=8.5, bold=True, color=MID_GRAY)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.paragraph_format.space_before = Pt(0)
        fr = fp.add_run("Versión 1.0  ·  Página ")
        font_run(fr, size=8.5, color=MID_GRAY)
        add_field(fp, "PAGE", "1")


def add_title(doc: Document, text: str, size=28, color=NAVY, after=8) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    font_run(run, size=size, bold=True, color=color)


def add_subtitle(doc: Document, text: str, size=14, color=MID_GRAY, after=18) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    font_run(run, size=size, color=color)


def add_para(doc: Document, text: str, *, bold_prefix: str | None = None, italic=False) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        font_run(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        font_run(r2, italic=italic)
    else:
        r = p.add_run(text)
        font_run(r, italic=italic)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.18
        r = p.add_run(item)
        font_run(r, size=10.5)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        font_run(r, size=10.5)


def add_callout(doc: Document, label: str, text: str, fill=PALE_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{label}: ")
    font_run(r, bold=True, color=NAVY)
    r2 = p.add_run(text)
    font_run(r2, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], font_size=9.2) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, PALE_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        font_run(r, size=9.2, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            font_run(r, size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def font(size: int, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def arrow(draw, p1, p2, color="#667085", width=4):
    draw.line([p1, p2], fill=color, width=width)
    x2, y2 = p2
    if abs(p2[0] - p1[0]) >= abs(p2[1] - p1[1]):
        pts = [(x2, y2), (x2 - 14 if x2 > p1[0] else x2 + 14, y2 - 8), (x2 - 14 if x2 > p1[0] else x2 + 14, y2 + 8)]
    else:
        pts = [(x2, y2), (x2 - 8, y2 - 14 if y2 > p1[1] else y2 + 14), (x2 + 8, y2 - 14 if y2 > p1[1] else y2 + 14)]
    draw.polygon(pts, fill=color)


def box(draw, xy, title, detail="", fill="#E8EEF5", outline="#2E74B5"):
    draw.rounded_rectangle(xy, radius=16, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    title_font = font(27, True)
    detail_font = font(20)
    tb = draw.textbbox((0, 0), title, font=title_font)
    draw.text(((x1 + x2 - (tb[2]-tb[0]))/2, y1+18), title, font=title_font, fill="#17365D")
    if detail:
        lines = detail.split("\n")
        y = y1 + 60
        for line in lines:
            db = draw.textbbox((0, 0), line, font=detail_font)
            draw.text(((x1+x2-(db[2]-db[0]))/2, y), line, font=detail_font, fill="#344054")
            y += 26


def make_diagrams() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    paths = {}
    img = Image.new("RGB", (1500, 850), "white")
    d = ImageDraw.Draw(img)
    box(d, (45, 300, 300, 500), "Usuario", "Navegador web")
    box(d, (440, 260, 770, 540), "Aplicación RAG", "Gradio · LangChain\nFly.io / Docker")
    box(d, (930, 80, 1435, 280), "OpenAI API", "Embeddings · generación\nre-ranking estructurado")
    box(d, (930, 350, 1435, 550), "Qdrant Cloud", "Colección vectorial\n192 chunks legales")
    box(d, (930, 620, 1435, 790), "Fly.io", "HTTPS · health check\n1 VM · región gru")
    arrow(d, (300, 400), (440, 400)); arrow(d, (770, 340), (930, 180)); arrow(d, (770, 430), (930, 450)); arrow(d, (770, 500), (930, 690))
    p = ASSET_DIR / "contexto_arquitectura.png"; img.save(p); paths["context"] = p

    img = Image.new("RGB", (1500, 900), "white"); d = ImageDraw.Draw(img)
    stages = [
        ("1. Pregunta", "texto del usuario"), ("2. Embedding", "text-embedding-3-large\n256 dimensiones"),
        ("3. Recuperación", "Qdrant · top 20"), ("4. Re-ranking", "LLM · score 0..1\numbral 0,55"),
        ("5. Generación", "contexto filtrado\nrespuesta + cita"), ("6. Entrega", "Markdown en Gradio"),
    ]
    x_positions = [40, 290, 540, 790, 1040, 1290]
    for i, (title, detail) in enumerate(stages):
        box(d, (x_positions[i], 290, x_positions[i]+190, 570), title, detail, fill="#F2F4F7")
        if i < len(stages)-1: arrow(d, (x_positions[i]+190, 430), (x_positions[i+1], 430), width=3)
    p = ASSET_DIR / "flujo_consulta.png"; img.save(p); paths["flow"] = p

    img = Image.new("RGB", (1500, 900), "white"); d = ImageDraw.Draw(img)
    box(d, (520, 45, 980, 190), "Repositorio GitHub", "main · Dockerfile · fly.toml")
    box(d, (520, 300, 980, 500), "Fly.io Machine", "Python 3.11 slim · usuario app\npuerto interno 7860 · HTTPS")
    box(d, (55, 650, 500, 840), "OpenAI", "API externa · secretos Fly")
    box(d, (1000, 650, 1445, 840), "Qdrant Cloud", "ley_21719_rag_prod\n256D · cosine")
    arrow(d, (750, 190), (750, 300)); arrow(d, (610, 500), (360, 650)); arrow(d, (890, 500), (1220, 650))
    p = ASSET_DIR / "despliegue.png"; img.save(p); paths["deploy"] = p
    return paths


def add_figure(doc: Document, path: Path, caption: str, width=6.55) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    cp = doc.add_paragraph(style="Caption")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.add_run(caption)


def build() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = make_diagrams()
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)

    # Portada — memo_masthead sobre compact_reference_guide.
    doc.add_paragraph().paragraph_format.space_after = Pt(48)
    kicker = doc.add_paragraph()
    r = kicker.add_run("SOFTWARE DESIGN DESCRIPTION")
    font_run(r, size=11, bold=True, color=BLUE)
    kicker.paragraph_format.space_after = Pt(10)
    add_title(doc, "Sistema RAG productivo\nLey 19.628 de Chile", size=28, after=8)
    add_subtitle(doc, "Descripción detallada del diseño conforme a IEEE 1016-2009", size=15, after=28)
    add_table(doc, ["Control", "Valor"], [
        ["Identificador", "SDD-RAG-L19628-001"],
        ["Versión", "1.0"],
        ["Estado", "Línea base de producción"],
        ["Fecha", "1 de agosto de 2026"],
        ["Sistema", "Asistente RAG · Ley 19.628"],
        ["URL productiva", "https://rag-ley-19628-grupo1.fly.dev"],
        ["Repositorio", "github.com/luiscaroca-ai/rag-ley-19628"],
        ["Estándar de organización", "IEEE 1016-2009"],
    ], [2100, 7260], font_size=10)
    add_callout(doc, "Clasificación", "Documento técnico de diseño. No contiene secretos, claves API ni datos personales. La aplicación no reemplaza asesoría jurídica.", fill=LIGHT_GRAY)
    page_break(doc)

    doc.add_heading("Control documental", level=1)
    add_table(doc, ["Versión", "Fecha", "Descripción", "Responsable"], [
        ["1.0", "2026-08-01", "Emisión inicial basada en la versión productiva desplegada.", "Equipo Grupo 1"],
    ], [1100, 1500, 5160, 1600])
    doc.add_heading("Aprobaciones", level=2)
    add_table(doc, ["Rol", "Nombre", "Estado", "Fecha"], [
        ["Propietario del producto", "Grupo 1", "Pendiente de firma", "—"],
        ["Responsable técnico", "Grupo 1", "Pendiente de firma", "—"],
        ["Revisor de seguridad", "Por designar", "Pendiente", "—"],
    ], [2400, 2600, 2500, 1860])
    doc.add_heading("Declaración de conformidad", level=2)
    add_para(doc, "Este SDD organiza la información mediante vistas regidas por puntos de vista y documenta interesados, preocupaciones, entidades, relaciones, atributos, restricciones y fundamentos de diseño. Se emplean los doce puntos de vista definidos por IEEE 1016-2009 cuando resultan aplicables. IEEE clasifica actualmente el estándar como Inactive-Reserved; la conformidad declarada se refiere a su estructura y contenido, no a una certificación de IEEE.")
    page_break(doc)

    doc.add_heading("Contenido", level=1)
    contents = [
        "1. Introducción y alcance", "2. Referencias y definiciones", "3. Interesados, preocupaciones y trazabilidad",
        "4. Vista de contexto", "5. Vista de composición", "6. Vista lógica", "7. Vista de dependencias",
        "8. Vista de información", "9. Vista de uso de patrones", "10. Vista de interfaces", "11. Vista de estructura",
        "12. Vista de interacción", "13. Vista de dinámica de estados", "14. Vista de algoritmos", "15. Vista de recursos",
        "16. Decisiones y fundamentos de diseño", "17. Seguridad, privacidad y operación", "18. Verificación y criterios de aceptación",
        "19. Riesgos, deuda técnica y evolución", "Anexos A–D"
    ]
    add_numbered(doc, contents)
    add_callout(doc, "Navegación", "Los identificadores VP-xx corresponden a puntos de vista; DE-xx a entidades; INT-xx a interfaces; AD-xx a decisiones; y RQ-xx a requisitos de diseño.")
    page_break(doc)

    doc.add_heading("1. Introducción y alcance", level=1)
    doc.add_heading("1.1 Propósito", level=2)
    add_para(doc, "Registrar y comunicar el diseño de la aplicación productiva que responde consultas sobre la Ley 19.628 de protección de datos personales de Chile usando recuperación aumentada por generación (RAG), fuentes legales y re-ranking. El documento sirve a desarrollo, pruebas, operación, seguridad, mantenimiento y evaluación académica.")
    doc.add_heading("1.2 Alcance del sujeto de diseño", level=2)
    add_bullets(doc, [
        "Incluye la aplicación web Gradio, el pipeline LangChain, el chunking por artículo, la indexación, las integraciones OpenAI/Qdrant y el despliegue Docker/Fly.io.",
        "Incluye configuración, secretos, observabilidad básica, recuperación, re-ranking, generación y presentación de citas.",
        "Excluye la infraestructura interna administrada por OpenAI, Qdrant Cloud y Fly.io; se documentan solo sus contratos externos.",
        "Excluye asesoría jurídica, control de identidad, facturación, administración multiusuario y edición del corpus desde la interfaz.",
    ])
    doc.add_heading("1.3 Objetivos de diseño", level=2)
    add_table(doc, ["ID", "Objetivo", "Indicador de diseño"], [
        ["RQ-01", "Responder solo con evidencia recuperada.", "Prompt restrictivo y respuesta negativa explícita."],
        ["RQ-02", "Citar el artículo fuente.", "Metadata `articulo` incorporada al contexto."],
        ["RQ-03", "Mejorar relevancia en preguntas complejas.", "Top-20, juez LLM y umbral 0,55."],
        ["RQ-04", "Desplegar de forma reproducible.", "Dockerfile, fly.toml y configuración por entorno."],
        ["RQ-05", "Evitar exposición de secretos.", "`.env` ignorado y secretos cifrados en Fly.io."],
        ["RQ-06", "Mantener continuidad operativa simple.", "Health check, auto-start y una máquina siempre activa."],
    ], [900, 3980, 4480])
    doc.add_heading("1.4 Restricciones", level=2)
    add_bullets(doc, [
        "Corpus consolidado `Ley_19628_refundida_21719.txt`; cambios legales requieren reindexación.",
        "Embeddings de 256 dimensiones y distancia coseno; la colección debe tener exactamente esa configuración.",
        "Dependencia de conectividad externa y cuotas de OpenAI/Qdrant.",
        "Una sola máquina productiva mientras Gradio mantenga estado SSE en memoria y no exista afinidad de sesión.",
    ])

    page_break(doc)
    doc.add_heading("2. Referencias y definiciones", level=1)
    add_table(doc, ["Referencia", "Uso en este SDD"], [
        ["IEEE Std 1016-2009", "Organización del SDD mediante vistas, puntos de vista, preocupaciones y fundamentos."],
        ["Ley 19.628, texto consolidado con Ley 21.719", "Corpus jurídico del sistema."],
        ["Repositorio `luiscaroca-ai/rag-ley-19628`", "Fuente de verdad del diseño implementado."],
        ["Documentación OpenAI API", "Contratos de embeddings, chat y salida estructurada."],
        ["Documentación Qdrant", "Colecciones vectoriales, similitud coseno y búsqueda."],
        ["Documentación Fly.io", "Máquinas, secretos, health checks, HTTPS y despliegue."],
    ], [3000, 6360])
    doc.add_heading("2.1 Términos", level=2)
    add_table(doc, ["Término", "Definición operativa"], [
        ["SDD", "Software Design Description: representación del diseño para registrar y comunicar información."],
        ["RAG", "Generación aumentada por recuperación; el LLM responde usando pasajes recuperados."],
        ["Chunk", "Unidad indexable de texto legal con metadata de fuente, artículo y parte."],
        ["Embedding", "Vector numérico que representa semánticamente una consulta o un chunk."],
        ["Re-ranking", "Segunda evaluación de relevancia aplicada a candidatos recuperados."],
        ["Umbral", "Score mínimo 0,55 para que un candidato pase al contexto final."],
        ["Vista", "Representación del diseño que atiende preocupaciones específicas."],
        ["Punto de vista", "Convenciones y elementos usados para construir una vista."],
    ], [2000, 7360])

    page_break(doc)
    doc.add_heading("3. Interesados, preocupaciones y trazabilidad", level=1)
    add_table(doc, ["Interesado", "Preocupaciones principales", "Vistas"], [
        ["Usuario final", "Exactitud, claridad, citas, disponibilidad, privacidad.", "VP-01, VP-07, VP-09"],
        ["Docente/evaluador", "Justificación técnica, trazabilidad, evidencia de RAG.", "VP-02, VP-06, VP-10"],
        ["Desarrollador", "Modularidad, contratos, algoritmos, manejo de errores.", "VP-03, VP-08, VP-11"],
        ["Operaciones", "Despliegue, salud, recursos, rollback, costos.", "VP-04, VP-12"],
        ["Seguridad", "Secretos, exposición pública, proveedores, abuso.", "VP-04, VP-07, VP-12"],
        ["Mantenedor legal", "Versión del corpus, citas, reindexación y vigencia.", "VP-05, VP-10"],
    ], [1900, 4840, 2620])
    doc.add_heading("3.1 Matriz de cobertura IEEE 1016", level=2)
    viewpoints = [
        ("VP-01", "Contexto", "Servicios, actores y límites", "§4"), ("VP-02", "Composición", "Subsistemas y ensamblaje", "§5"),
        ("VP-03", "Lógico", "Responsabilidades y abstracciones", "§6"), ("VP-04", "Dependencias", "Acceso, ejecución e impacto", "§7"),
        ("VP-05", "Información", "Datos persistentes y distribución", "§8"), ("VP-06", "Patrones", "Soluciones reutilizadas", "§9"),
        ("VP-07", "Interfaces", "Contratos internos y externos", "§10"), ("VP-08", "Estructura", "Clases, módulos y relaciones", "§11"),
        ("VP-09", "Interacción", "Mensajes y secuencias", "§12"), ("VP-10", "Estados", "Estados y transiciones", "§13"),
        ("VP-11", "Algoritmos", "Procedimientos y reglas", "§14"), ("VP-12", "Recursos", "Cómputo, red, almacenamiento", "§15"),
    ]
    add_table(doc, ["ID", "Punto de vista", "Preocupación", "Ubicación"], viewpoints, [900, 1900, 4960, 1600])

    page_break(doc)
    doc.add_heading("4. Vista de contexto — VP-01", level=1)
    add_para(doc, "Perspectiva de caja negra que identifica servicios, actores y sistemas externos. El sujeto de diseño presta un único servicio público: responder preguntas legales con evidencia del corpus indexado.")
    add_figure(doc, diagrams["context"], "Figura 1 — Contexto y fronteras de confianza de la aplicación productiva.")
    add_table(doc, ["Actor/sistema", "Responsabilidad", "Flujo intercambiado"], [
        ["Usuario", "Formula preguntas y consume respuestas.", "Pregunta UTF-8; respuesta Markdown."],
        ["Fly.io", "Ejecuta el contenedor y termina HTTPS.", "HTTP/S, health checks y secretos."],
        ["OpenAI API", "Genera embeddings, scores y respuesta.", "Consultas, pasajes y contexto; resultados de modelo."],
        ["Qdrant Cloud", "Persiste vectores y recupera candidatos.", "Vector de consulta; documentos y metadata."],
        ["GitHub", "Versiona código y configuración.", "Fuentes sin secretos ni corpus local."],
    ], [2200, 3300, 3860])
    add_callout(doc, "Frontera de confianza", "La pregunta y los pasajes recuperados salen del contenedor hacia OpenAI. Las claves se inyectan como secretos; nunca se presentan al usuario ni se registran deliberadamente.", fill="FFF4E5")

    page_break(doc)
    doc.add_heading("5. Vista de composición — VP-02", level=1)
    add_table(doc, ["Entidad", "Tipo", "Responsabilidad", "Implementación"], [
        ["DE-01 UI web", "Componente", "Captura pregunta y presenta Markdown.", "`app.py` / Gradio"],
        ["DE-02 LegalRAG", "Componente", "Orquesta recuperación, re-ranking y generación.", "`pipeline.py`"],
        ["DE-03 Configuración", "Componente", "Resuelve variables y valida credenciales.", "`config.py`"],
        ["DE-04 Chunking", "Componente", "Segmenta por artículo y crea metadata.", "`chunking.py`"],
        ["DE-05 Indexador", "Proceso batch", "Crea la colección y carga documentos.", "`index_law.py`"],
        ["DE-06 Vector store", "Adaptador", "Conecta LangChain con Qdrant.", "QdrantVectorStore"],
        ["DE-07 Modelos", "Servicios", "Embeddings, evaluación y generación.", "OpenAIEmbeddings/ChatOpenAI"],
        ["DE-08 Plataforma", "Despliegue", "Ejecuta y expone el servicio.", "Docker + Fly.io"],
    ], [1100, 1500, 3820, 2940])
    doc.add_heading("5.1 Ensamblaje en tiempo de ejecución", level=2)
    add_numbered(doc, [
        "Fly inicia el contenedor y ejecuta `python app.py` como usuario no privilegiado `app`.",
        "La importación crea `LegalRAG`, valida credenciales y construye clientes OpenAI/Qdrant.",
        "Gradio publica el endpoint en `0.0.0.0:7860`; Fly Proxy expone 80/443 y fuerza HTTPS.",
        "Cada solicitud invoca `LegalRAG.answer`; no se persiste estado de conversación.",
    ])

    page_break(doc)
    doc.add_heading("6. Vista lógica — VP-03", level=1)
    add_para(doc, "La descomposición lógica separa presentación, aplicación, dominio documental e infraestructura. La separación es parcial: `LegalRAG` concentra orquestación y adaptadores de proveedores, una decisión aceptable para el alcance actual pero relevante para evolución.")
    add_table(doc, ["Capa", "Elementos", "Reglas"], [
        ["Presentación", "Gradio Interface", "No contiene lógica de recuperación; valida pregunta vacía indirectamente."],
        ["Aplicación", "LegalRAG", "Coordina casos de uso; retorna texto listo para presentar."],
        ["Dominio documental", "Document, chunk_law, metadata", "Artículo y fuente deben preservarse en todos los chunks."],
        ["Infraestructura", "OpenAI, Qdrant, dotenv, Fly", "Acceso por adaptadores/configuración; secretos fuera del código."],
    ], [1700, 3000, 4660])
    doc.add_heading("6.1 Invariantes", level=2)
    add_bullets(doc, [
        "Todo documento recuperable posee `source`, `articulo` y `parte`.",
        "La dimensión del embedding coincide con la colección: 256.",
        "La generación usa únicamente documentos que sobreviven al umbral cuando re-ranking está activo.",
        "Si no hay sobrevivientes, la salida exacta es «No se encuentra en el documento.».",
        "La aplicación no almacena sesión, historial ni identidad del usuario.",
    ])

    page_break(doc)
    doc.add_heading("7. Vista de dependencias — VP-04", level=1)
    add_table(doc, ["Origen", "Dependencia", "Tipo", "Efecto de falla", "Mitigación actual"], [
        ["UI", "LegalRAG", "Invocación", "No se responde.", "Mensaje de error de plataforma; pendiente manejo amigable."],
        ["LegalRAG", "OpenAI", "HTTPS síncrono", "Embedding, score o respuesta fallan.", "Timeout del SDK; health no cubre proveedor."],
        ["LegalRAG", "Qdrant", "HTTPS síncrono", "No hay recuperación.", "Colección productiva separada."],
        ["Aplicación", "Fly.io", "Runtime", "Servicio no disponible.", "Auto-start y health check."],
        ["Build", "PyPI/Docker Hub", "Construcción", "No se genera imagen.", "Imagen base versionada por tag; caché remota."],
        ["Indexador", "Corpus legal", "Archivo", "Índice incompleto o desactualizado.", "Nombre fijo y error si falta."],
    ], [1350, 1700, 1550, 2360, 2400], font_size=8.7)
    doc.add_heading("7.1 Orden de ejecución", level=2)
    add_numbered(doc, ["Configurar secretos.", "Preparar corpus.", "Ejecutar indexación productiva.", "Desplegar aplicación.", "Escalar a una máquina.", "Verificar health y consulta sintética."])

    page_break(doc)
    doc.add_heading("8. Vista de información — VP-05", level=1)
    add_table(doc, ["Entidad de datos", "Esquema", "Persistencia", "Sensibilidad"], [
        ["Corpus", "Texto UTF-8 consolidado", "Archivo local solo para indexación", "Público/legal"],
        ["Chunk", "page_content + metadata", "Payload Qdrant", "Público/legal"],
        ["Vector", "float[256]", "Qdrant `ley_21719_rag_prod`", "Derivado del corpus"],
        ["Pregunta", "string", "No persistida por la app", "Puede contener datos ingresados"],
        ["Score", "float 0..1", "Memoria de solicitud", "No sensible"],
        ["Respuesta", "Markdown string", "No persistida por la app", "Derivada de pregunta/contexto"],
        ["Secretos", "claves/URL", "Fly Secrets y `.env` local", "Confidencial"],
    ], [1700, 2600, 2960, 2100])
    doc.add_heading("8.1 Esquema de payload vectorial", level=2)
    add_table(doc, ["Campo", "Tipo", "Ejemplo", "Regla"], [
        ["page_content", "string", "Artículo 12...", "Texto indexado y recuperado."],
        ["metadata.source", "string", "Ley 19.628 (ref. 21.719)", "Valor de procedencia estable."],
        ["metadata.articulo", "string", "Artículo 12", "Base para cita visible."],
        ["metadata.parte", "string", "2/4", "Orden relativo dentro del artículo."],
        ["vector", "float[256]", "[0.012, …]", "Distancia COSINE."],
    ], [1800, 1500, 2960, 3100])
    add_callout(doc, "Retención", "La aplicación no implementa una base de datos de usuarios ni historial. Los proveedores externos pueden aplicar sus propias políticas de registro y retención; deben revisarse contractualmente antes de tratar información sensible.", fill="FFF4E5")

    page_break(doc)
    doc.add_heading("9. Vista de uso de patrones — VP-06", level=1)
    add_table(doc, ["Patrón", "Aplicación", "Beneficio", "Trade-off"], [
        ["RAG", "Recuperar evidencia antes de generar.", "Reduce alucinación y aporta citas.", "Depende de calidad del retrieval."],
        ["Pipeline", "Pregunta → embedding → retrieval → rerank → respuesta.", "Secuencia explícita y testeable.", "Latencia acumulada."],
        ["Adapter", "LangChain encapsula OpenAI/Qdrant.", "Menor acoplamiento a APIs crudas.", "Acoplamiento a versiones LangChain."],
        ["Configuration by environment", "Variables `.env`/Fly.", "Misma imagen entre entornos.", "Errores aparecen en arranque."],
        ["Contextual chunking", "Encabezado se antepone a fragmentos.", "Preserva sentido legal.", "Duplica texto y tokens."],
        ["Threshold gate", "Score mínimo 0,55.", "Evita contexto irrelevante.", "Puede producir falsos negativos."],
    ], [1800, 3150, 2500, 1910], font_size=8.8)
    doc.add_heading("9.1 Patrones descartados", level=2)
    add_bullets(doc, [
        "Conversación con memoria: descartada para reducir exposición de datos y complejidad de sesión.",
        "Múltiples máquinas activas: descartadas porque Gradio mantiene eventos SSE en memoria sin afinidad externa.",
        "Vector store local persistente: descartado para separar cómputo efímero de datos vectoriales.",
    ])

    page_break(doc)
    doc.add_heading("10. Vista de interfaces — VP-07", level=1)
    add_table(doc, ["ID", "Interfaz", "Entrada", "Salida/contrato", "Errores relevantes"], [
        ["INT-01", "Gradio `/answer`", "question:string", "Markdown:string", "Pregunta vacía; excepción upstream."],
        ["INT-02", "LegalRAG.answer", "question, rerank=True", "Respuesta o negativa exacta", "Credenciales/colección/modelo."],
        ["INT-03", "Qdrant search", "Vector/consulta, k", "Document[] con metadata", "Colección ausente; red; auth."],
        ["INT-04", "OpenAI embeddings", "Texto", "float[256]", "Cuota; auth; timeout."],
        ["INT-05", "OpenAI scorer", "Pregunta + pasaje", "Relevance.score", "Salida estructurada inválida."],
        ["INT-06", "OpenAI generator", "Pregunta + contexto", "Texto con cita", "Modelo no disponible; cuota."],
        ["INT-07", "Fly health", "GET `/`", "HTTP 200", "Arranque lento o proceso caído."],
    ], [900, 1700, 1750, 2700, 2310], font_size=8.4)
    doc.add_heading("10.1 Interfaz de usuario", level=2)
    add_bullets(doc, [
        "Campo de texto con ejemplos orientativos; salida Markdown etiquetada «Respuesta con fuente».",
        "No se solicita autenticación ni se almacenan datos del usuario.",
        "La respuesta debe citar el artículo entre corchetes; la UI no valida automáticamente la cita.",
        "URL pública HTTPS: `https://rag-ley-19628-grupo1.fly.dev`.",
    ])

    page_break(doc)
    doc.add_heading("11. Vista de estructura — VP-08", level=1)
    add_table(doc, ["Módulo/clase", "Miembros clave", "Relaciones"], [
        ["Settings", "source_file, collection, models, k, threshold", "Instancia singleton `settings`; consumida por indexador y LegalRAG."],
        ["LegalRAG", "embeddings, client, store, answer_chain, scorer", "Compone proveedores y Settings."],
        ["Relevance", "score: float [0,1]", "Modelo Pydantic de salida del scorer."],
        ["chunk_law", "text, chunk_size, overlap, limit", "Produce `list[Document]`."],
        ["format_context", "documents", "Transforma Documents a bloque con citas."],
        ["Gradio Interface", "fn=rag.answer", "Adaptador UI sobre LegalRAG."],
    ], [2200, 3560, 3600])
    doc.add_heading("11.1 Estructura del repositorio", level=2)
    add_para(doc, "`app.py` contiene presentación; `src/rag_ley/` contiene configuración, chunking y pipeline; `scripts/` contiene el proceso batch; `tests/` contiene pruebas; Dockerfile/compose/fly.toml describen ejecución; `notebooks/` conserva el origen exploratorio.")
    doc.add_heading("11.2 Cohesión y acoplamiento", level=2)
    add_bullets(doc, [
        "Alta cohesión en chunking y configuración.",
        "LegalRAG tiene responsabilidad amplia: construcción, retrieval, scoring y generación.",
        "Los proveedores están acoplados a clases LangChain concretas; una futura inversión de dependencias facilitaría pruebas unitarias.",
    ])

    page_break(doc)
    doc.add_heading("12. Vista de interacción — VP-09", level=1)
    add_figure(doc, diagrams["flow"], "Figura 2 — Secuencia lógica de una consulta con re-ranking.")
    add_table(doc, ["Paso", "Emisor → receptor", "Mensaje", "Resultado"], [
        ["1", "Usuario → Gradio", "Pregunta", "Evento de consulta."],
        ["2", "LegalRAG → OpenAI", "Embedding de pregunta", "Vector 256D."],
        ["3", "LegalRAG → Qdrant", "similarity_search(k=20)", "Candidatos."],
        ["4", "LegalRAG → OpenAI", "Batch pregunta/pasaje", "Scores de relevancia."],
        ["5", "LegalRAG", "Ordenar y filtrar ≥0,55", "Sobrevivientes."],
        ["6", "LegalRAG → OpenAI", "Contexto + pregunta", "Respuesta citada."],
        ["7", "Gradio → Usuario", "Markdown", "Visualización final."],
    ], [900, 2400, 3300, 2760])
    add_callout(doc, "Latencia", "Una consulta re-ranked realiza una búsqueda vectorial, un batch de puntuación y una generación. La latencia depende principalmente de OpenAI y del número de candidatos.")

    page_break(doc)
    doc.add_heading("13. Vista de dinámica de estados — VP-10", level=1)
    add_table(doc, ["Estado", "Evento de entrada", "Acción", "Siguiente estado"], [
        ["Detenida", "Solicitud pública", "Fly auto-start", "Iniciando"],
        ["Iniciando", "Contenedor ejecuta app.py", "Validar credenciales y conectar proveedores", "Lista o Fallida"],
        ["Lista", "Pregunta vacía", "Retornar instrucción", "Lista"],
        ["Lista", "Pregunta válida", "Ejecutar pipeline", "Procesando"],
        ["Procesando", "Sin sobrevivientes", "Retornar negativa", "Lista"],
        ["Procesando", "Respuesta generada", "Entregar Markdown", "Lista"],
        ["Procesando", "Excepción no controlada", "HTTP/error de evento", "Lista o Fallida"],
        ["Lista", "Inactividad/exceso de capacidad", "Auto-stop según Fly", "Detenida"],
        ["Cualquier", "Proceso termina", "Health falla/reinicio", "Iniciando o Fallida"],
    ], [1600, 2300, 3100, 2360], font_size=8.8)
    doc.add_heading("13.1 Estado de la consulta", level=2)
    add_para(doc, "La consulta es efímera y no se comparte entre solicitudes. Gradio conserva temporalmente el evento SSE en memoria; por ello el despliegue se fija en una máquina. No existe recuperación de una consulta en curso ante reinicio.")

    page_break(doc)
    doc.add_heading("14. Vista de algoritmos — VP-11", level=1)
    doc.add_heading("14.1 Chunking estructural", level=2)
    add_numbered(doc, [
        "Dividir el corpus mediante expresión regular con look-ahead en encabezados «Artículo». Complejidad O(n) aproximada sobre el texto.",
        "Para cada bloque, extraer etiqueta y encabezado.",
        "Conservar artículos de hasta 1.500 caracteres; dividir los mayores en piezas de 1.200 con solape 150.",
        "Anteponer el encabezado a las piezas que no comienzan por «Artículo».",
        "Emitir Document con source, articulo y parte. Resultado actual: 192 chunks.",
    ])
    doc.add_heading("14.2 Indexación", level=2)
    add_numbered(doc, [
        "Validar credenciales y existencia del corpus.", "Generar chunks.", "Crear cliente Qdrant y embeddings.",
        "Si la colección existe, eliminarla.", "Crear colección COSINE de 256 dimensiones.", "Calcular y cargar embeddings." ])
    add_callout(doc, "Advertencia", "La indexación es destructiva para la colección objetivo. Producción usa `ley_21719_rag_prod` para aislarla de desarrollo.", fill="FDECEC")
    doc.add_heading("14.3 Recuperación y re-ranking", level=2)
    add_para(doc, "Recuperar k=20 documentos por similitud; puntuar cada uno con un LLM de salida estructurada; ordenar de mayor a menor; conservar score ≥0,55. Si el conjunto queda vacío, terminar sin generación. En caso contrario, concatenar contexto y solicitar respuesta al modelo generador.")
    doc.add_heading("14.4 Reglas del prompt", level=2)
    add_bullets(doc, ["Usar exclusivamente el contexto.", "No inventar información.", "Ser conciso.", "Citar el artículo entre corchetes.", "Usar negativa exacta cuando no haya evidencia."])

    page_break(doc)
    doc.add_heading("15. Vista de recursos — VP-12", level=1)
    add_figure(doc, diagrams["deploy"], "Figura 3 — Topología de despliegue productivo.", width=4.5)
    add_table(doc, ["Recurso", "Asignación", "Uso", "Límite/observación"], [
        ["Fly Machine", "shared-cpu-1x, 1 GB", "Gradio + pipeline", "Una instancia por afinidad SSE."],
        ["Red", "HTTPS público", "UI y APIs externas", "Fly fuerza HTTPS."],
        ["Puerto", "7860 interno", "Servidor Gradio", "Expuesto por Fly Proxy."],
        ["Qdrant", "Colección 256D COSINE", "192 vectores + payload", "Servicio externo."],
        ["OpenAI", "3 modelos configurables", "Embedding, score, generación", "Costo/cuotas variables."],
        ["Secretos", "Fly Secrets", "Tres credenciales", "No incluidos en imagen."],
        ["Health", "GET / cada 30 s", "Disponibilidad del proceso", "No prueba OpenAI/Qdrant."],
    ], [1700, 2200, 2700, 2760], font_size=8.8)
    doc.add_heading("15.1 Parámetros productivos", level=2)
    add_table(doc, ["Parámetro", "Valor"], [
        ["App", "rag-ley-19628-grupo1"], ["Región primaria", "gru (São Paulo)"], ["HTTPS", "forzado"],
        ["Auto-start", "habilitado"], ["Auto-stop", "stop"], ["Mínimo activo", "1"],
        ["Concurrencia", "soft 20 / hard 30 solicitudes"], ["Estrategia", "rolling"],
    ], [3000, 6360])

    page_break(doc)
    doc.add_heading("16. Decisiones y fundamentos de diseño", level=1)
    add_table(doc, ["ID", "Decisión", "Justificación", "Alternativas/impacto"], [
        ["AD-01", "Chunking por artículo", "Alinea unidad semántica con estructura legal y citas.", "Fixed-size: mezcla o fragmenta normas."],
        ["AD-02", "Embeddings 256D", "Reduce almacenamiento/costo manteniendo modelo grande.", "Dimensión completa: mayor costo."],
        ["AD-03", "Qdrant Cloud", "Persistencia y búsqueda separadas del cómputo efímero.", "Local: requiere volumen y operación."],
        ["AD-04", "Re-ranking LLM", "Mejora preguntas específicas o multi-salto.", "Cross-encoder: menor dependencia LLM, no implementado."],
        ["AD-05", "Umbral 0,55", "Filtra evidencia débil y habilita negativa.", "Valor requiere calibración continua."],
        ["AD-06", "Una Fly Machine", "Evita pérdida de eventos SSE entre instancias.", "HA requiere sticky sessions/cola externa."],
        ["AD-07", "Colección prod separada", "Evita que experimentos destruyan producción.", "Colección compartida: menor costo, mayor riesgo."],
        ["AD-08", "Sin memoria conversacional", "Minimiza privacidad y complejidad.", "Memoria: mejor UX, más retención y estado."],
    ], [900, 2100, 3460, 2900], font_size=8.3)

    page_break(doc)
    doc.add_heading("17. Seguridad, privacidad y operación", level=1)
    doc.add_heading("17.1 Controles implementados", level=2)
    add_bullets(doc, [
        "Contenedor ejecutado como usuario del sistema no privilegiado.", "`.env` excluido del repositorio y del contexto Docker.",
        "Secretos cargados en Fly.io y no codificados en `fly.toml`.", "HTTPS obligatorio en el borde Fly.",
        "Sin persistencia de usuarios, sesiones o historial en la aplicación.", "Colección productiva aislada por nombre.",
    ])
    doc.add_heading("17.2 Amenazas y mitigaciones", level=2)
    add_table(doc, ["Amenaza", "Impacto", "Control actual", "Recomendación"], [
        ["Abuso público/costos", "Consumo de tokens", "Límite de concurrencia", "Añadir rate limiting y presupuesto."],
        ["Prompt injection", "Respuesta fuera de alcance", "Prompt restrictivo", "Filtrar entradas y evaluar ataques."],
        ["Exfiltración de secretos", "Compromiso de proveedores", "Fly Secrets/no root", "Rotación y scopes mínimos."],
        ["Corpus alterado", "Respuestas incorrectas", "Fuente local controlada", "Hash, firma y revisión de cambios."],
        ["Dependencia de terceros", "Indisponibilidad", "Health de proceso", "Circuit breaker y estado degradado."],
        ["Datos personales en preguntas", "Privacidad", "No persistencia local", "Aviso de privacidad y revisión de retención externa."],
    ], [1900, 1900, 2600, 2960], font_size=8.5)
    doc.add_heading("17.3 Operación", level=2)
    add_bullets(doc, [
        "Despliegue: `flyctl deploy`; mantener `flyctl scale count 1 --yes`.",
        "Observación: `flyctl status`, `flyctl checks list`, `flyctl logs` y consulta sintética.",
        "Rollback: seleccionar release/imagen anterior en Fly.io; verificar compatibilidad con la colección.",
        "Reindexación: ejecutar el perfil `index` contra `ley_21719_rag_prod` durante ventana controlada.",
    ])

    page_break(doc)
    doc.add_heading("18. Verificación y criterios de aceptación", level=1)
    add_table(doc, ["ID", "Criterio", "Método", "Resultado base"], [
        ["AC-01", "La imagen Docker se construye.", "`docker compose build`", "Aprobado"],
        ["AC-02", "Chunking preserva artículo.", "pytest", "1 prueba aprobada"],
        ["AC-03", "Fly config es válida.", "`flyctl config validate`", "Aprobado"],
        ["AC-04", "La máquina está saludable.", "Fly status/check", "1/1 passing"],
        ["AC-05", "La URL responde por HTTPS.", "HTTP GET", "200"],
        ["AC-06", "RAG responde con cita.", "Consulta sintética", "Artículos 2, 12 y 30 bis/51 verificados"],
        ["AC-07", "No se publican secretos.", "gitignore/dockerignore/revisión", "Aprobado"],
        ["AC-08", "Producción usa colección separada.", "fly.toml/Qdrant", "ley_21719_rag_prod"],
    ], [1000, 3550, 2600, 2210])
    doc.add_heading("18.1 Estrategia de pruebas recomendada", level=2)
    add_bullets(doc, [
        "Unitarias: regex, límites, metadata, formato de contexto y umbral.",
        "Integración: colección temporal Qdrant y adaptadores OpenAI simulados.",
        "Contrato: esquema Pydantic Relevance y dimensiones de embedding.",
        "End-to-end: preguntas de oro, pregunta trampa, citas y latencia.",
        "Seguridad: prompt injection, entradas extensas, abuso concurrente y escaneo de secretos.",
    ])

    page_break(doc)
    doc.add_heading("19. Riesgos, deuda técnica y evolución", level=1)
    add_table(doc, ["ID", "Riesgo/deuda", "Prioridad", "Tratamiento recomendado"], [
        ["RT-01", "Excepciones upstream no se transforman en mensajes amigables.", "Alta", "Manejo tipado, retry limitado y logging estructurado."],
        ["RT-02", "Health check solo valida la UI.", "Media", "Readiness separada para Qdrant/OpenAI con cache."],
        ["RT-03", "Una máquina limita alta disponibilidad.", "Alta", "API stateless o backend de eventos/afinidad."],
        ["RT-04", "Dependencias con rangos amplios.", "Media", "Lockfile y actualización automatizada."],
        ["RT-05", "Solo una prueba unitaria.", "Alta", "Cobertura de pipeline y pruebas contractuales."],
        ["RT-06", "El índice se reemplaza destructivamente.", "Alta", "Colecciones versionadas y alias atómico."],
        ["RT-07", "Sin autenticación ni rate limit.", "Alta", "Control de acceso o WAF/rate limiter."],
        ["RT-08", "Vigencia legal manual.", "Alta", "Proceso formal de actualización y hash de corpus."],
        ["RT-09", "Modelo y umbral sin telemetría.", "Media", "Evaluación continua y registro anonimizado de métricas."],
    ], [900, 4300, 1300, 2860], font_size=8.5)
    doc.add_heading("19.1 Hoja de ruta sugerida", level=2)
    add_numbered(doc, [
        "Corto plazo: manejo de errores, rate limit, pruebas y lockfile.",
        "Mediano plazo: índice versionado, observabilidad, evaluación automática y actualización legal controlada.",
        "Largo plazo: API stateless, alta disponibilidad, autenticación y gobernanza de datos.",
    ])

    page_break(doc)
    doc.add_heading("Anexo A — Matriz de trazabilidad", level=1)
    add_table(doc, ["Requisito", "Entidad", "Vista", "Decisión", "Aceptación"], [
        ["RQ-01", "DE-02", "VP-09/11", "AD-01/04/05", "AC-06"], ["RQ-02", "DE-04/02", "VP-05/07", "AD-01", "AC-06"],
        ["RQ-03", "DE-02/07", "VP-09/11", "AD-04/05", "AC-06"], ["RQ-04", "DE-08", "VP-02/12", "AD-06", "AC-01/03/04/05"],
        ["RQ-05", "DE-03/08", "VP-04/12", "AD-08", "AC-07"], ["RQ-06", "DE-08", "VP-10/12", "AD-06", "AC-04/05"],
    ], [1500, 1900, 1900, 2200, 1860])
    doc.add_heading("Anexo B — Variables de configuración", level=1)
    add_table(doc, ["Variable", "Obligatoria", "Valor no secreto/base", "Función"], [
        ["OPENAI_API_KEY", "Sí", "Secreto", "Autenticación OpenAI."], ["QDRANT_URL", "Sí", "Secreto/configuración", "Endpoint del clúster."],
        ["QDRANT_API_KEY", "Sí", "Secreto", "Autenticación Qdrant."], ["QDRANT_COLLECTION", "No", "ley_21719_rag_prod", "Colección de producción."],
        ["OPENAI_EMBED_MODEL", "No", "text-embedding-3-large", "Modelo de embedding."], ["OPENAI_EMBED_DIMS", "No", "256", "Dimensión vectorial."],
        ["OPENAI_GEN_MODEL", "No", "gpt-5.4-mini", "Generación."], ["OPENAI_RERANK_MODEL", "No", "gpt-5.4-nano", "Puntuación."],
        ["RAG_THRESHOLD", "No", "0.55", "Filtro de relevancia."],
    ], [2600, 1400, 2600, 2760], font_size=8.6)

    page_break(doc)
    doc.add_heading("Anexo C — Lista de comprobación operativa", level=1)
    add_bullets(doc, [
        "Confirmar versión y hash del corpus legal.", "Confirmar colección objetivo productiva.", "Respaldar o versionar colección anterior.",
        "Ejecutar indexación y validar conteo de chunks.", "Validar secretos Fly sin imprimir valores.", "Ejecutar `flyctl config validate`.",
        "Desplegar, reducir a una máquina y esperar health passing.", "Ejecutar preguntas de oro y pregunta trampa.",
        "Revisar logs por excepciones y latencia.", "Actualizar SDD, README y registro de cambios.",
    ])
    doc.add_heading("Anexo D — Bibliografía", level=1)
    refs = [
        "IEEE Standards Association. IEEE Std 1016-2009, IEEE Standard for Information Technology—Systems Design—Software Design Descriptions. Publicado el 20 de julio de 2009; estado actual Inactive-Reserved.",
        "Biblioteca del Congreso Nacional de Chile. Ley 19.628 sobre protección de la vida privada, texto consolidado con modificaciones de la Ley 21.719.",
        "OpenAI. Documentación oficial de API y modelos.",
        "Qdrant. Documentación oficial de Qdrant Client y colecciones vectoriales.",
        "Fly.io. Documentación oficial de Dockerfile deployment, Machines, secrets y app configuration.",
        "LangChain. Documentación de LangChain Core, OpenAI y Qdrant integrations.",
    ]
    add_numbered(doc, refs)
    add_callout(doc, "Fin del documento", "Este SDD describe la línea base productiva observada al 1 de agosto de 2026. Toda modificación arquitectónica relevante debe actualizar la versión, la matriz de trazabilidad y las decisiones de diseño.", fill=LIGHT_GRAY)

    # Propiedades y actualización de campos.
    doc.core_properties.title = "Software Design Description — RAG Ley 19.628"
    doc.core_properties.subject = "IEEE 1016-2009"
    doc.core_properties.author = "Grupo 1"
    doc.core_properties.keywords = "IEEE 1016, SDD, RAG, Ley 19.628, Fly.io, Qdrant, OpenAI"
    settings = doc.settings._element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
